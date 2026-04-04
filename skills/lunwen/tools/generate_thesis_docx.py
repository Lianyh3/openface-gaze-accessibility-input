from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

try:
    from latex2mathml.converter import convert as latex_to_mathml
except Exception:
    latex_to_mathml = None

SPECIAL_CENTERED_HEADINGS = {
    "摘要",
    "中文摘要",
    "abstract",
    "目录",
    "目　　录",
    "设计总说明",
    "introduction",
    "参考文献",
    "致谢",
    "鸣谢",
    "声明",
    "在学期间参加课题的研究成果",
}

ABSTRACT_HEADINGS = {"摘要", "中文摘要", "abstract"}
TOC_HEADINGS = {"目录", "目　　录"}
BACK_MATTER_HEADINGS = {"参考文献", "致谢", "鸣谢", "声明", "在学期间参加课题的研究成果"}
APPENDIX_PREFIXES = ("附录",)
FRONT_MATTER_ORDER_TAGS = ("cn_abstract", "en_abstract", "toc")

DEFAULT_REPLACE_MARKERS = [
    "设计总说明",
    "摘 要",
    "摘要",
    "abstract",
    "introduction",
    "1 绪论",
    "1 （第一层标题）",
    "（第一层标题）",
]
MML2OMML_XSL = Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL")


def default_style_spec() -> dict:
    return {
        "document_title": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 16,
            "bold": True,
            "alignment": "center",
            "line_spacing_pt": 20,
            "space_before_pt": 20,
            "space_after_pt": 20,
        },
        "centered_heading": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 15,
            "bold": True,
            "alignment": "center",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
            "page_break_before": True,
            "outline_level": 0,
        },
        "abstract_heading_en": {
            "east_asia_font": "Times New Roman",
            "latin_font": "Times New Roman",
            "size_pt": 15,
            "bold": True,
            "alignment": "center",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
            "page_break_before": True,
            "outline_level": 0,
        },
        "heading1": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 12,
            "bold": True,
            "alignment": "center",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
            "page_break_before": True,
            "outline_level": 0,
        },
        "heading2": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 12,
            "bold": True,
            "alignment": "left",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
            "outline_level": 1,
        },
        "heading3": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 12,
            "bold": True,
            "alignment": "left",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
            "left_indent_pt": 21,
            "outline_level": 2,
        },
        "heading4": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 12,
            "bold": True,
            "alignment": "left",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
            "left_indent_pt": 21,
            "outline_level": 3,
        },
        "body": {
            "east_asia_font": "宋体",
            "latin_font": "Times New Roman",
            "size_pt": 10.5,
            "bold": False,
            "alignment": "left",
            "line_spacing_pt": 20,
            "first_line_indent_pt": 21,
            "space_before_pt": 0,
            "space_after_pt": 0,
        },
        "abstract_body_en": {
            "east_asia_font": "Times New Roman",
            "latin_font": "Times New Roman",
            "size_pt": 10.5,
            "bold": False,
            "alignment": "left",
            "line_spacing_pt": 20,
            "first_line_indent_pt": 21,
            "space_before_pt": 0,
            "space_after_pt": 0,
        },
        "caption": {
            "east_asia_font": "宋体",
            "latin_font": "Times New Roman",
            "size_pt": 9,
            "bold": False,
            "alignment": "center",
            "line_spacing_rule": "single",
            "space_before_pt": 0,
            "space_after_pt": 0,
        },
        "table_text": {
            "east_asia_font": "宋体",
            "latin_font": "Times New Roman",
            "size_pt": 10.5,
            "bold": False,
            "alignment": "center",
            "line_spacing_pt": 20,
            "three_line_table": True,
        },
        "reference": {
            "east_asia_font": "宋体",
            "latin_font": "Times New Roman",
            "size_pt": 9,
            "bold": False,
            "alignment": "left",
            "line_spacing_pt": 16,
            "first_line_indent_pt": -21,
            "left_indent_pt": 21,
            "space_before_pt": 0,
            "space_after_pt": 0,
        },
        "keywords_cn": {
            "label_east_asia_font": "黑体",
            "label_latin_font": "Times New Roman",
            "content_east_asia_font": "宋体",
            "content_latin_font": "Times New Roman",
            "size_pt": 10.5,
            "alignment": "left",
            "line_spacing_pt": 20,
        },
        "keywords_en": {
            "label_east_asia_font": "Times New Roman",
            "label_latin_font": "Times New Roman",
            "content_east_asia_font": "Times New Roman",
            "content_latin_font": "Times New Roman",
            "size_pt": 10.5,
            "alignment": "left",
            "line_spacing_pt": 20,
        },
        "toc_heading": {
            "east_asia_font": "黑体",
            "latin_font": "Times New Roman",
            "size_pt": 15,
            "bold": True,
            "alignment": "center",
            "line_spacing_pt": 20,
            "space_before_pt": 10,
            "space_after_pt": 10,
        },
    }


def merge_style_spec(overrides: dict | None) -> dict:
    spec = default_style_spec()
    if not overrides:
        return spec
    for key, value in overrides.items():
        if isinstance(value, dict) and isinstance(spec.get(key), dict):
            spec[key].update(value)
        else:
            spec[key] = value
    return spec


def load_style_spec(path: Path | None) -> dict:
    if path is None or not path.exists():
        return default_style_spec()
    data = json.loads(path.read_text(encoding="utf-8"))
    if "style_spec" in data and isinstance(data["style_spec"], dict):
        data = data["style_spec"]
    return merge_style_spec(data)


def load_style_spec_payload(path: Path | None) -> tuple[dict, dict]:
    if path is None or not path.exists():
        return default_style_spec(), {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    styles = payload["style_spec"] if "style_spec" in payload and isinstance(payload["style_spec"], dict) else payload
    return merge_style_spec(styles), payload


def load_image_map(path: Path | None) -> dict[str, Path]:
    if path is None or not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    return {k: Path(v) for k, v in data.items()}


def build_math_transform():
    if latex_to_mathml is None or not MML2OMML_XSL.exists():
        return None
    xslt_root = etree.parse(str(MML2OMML_XSL))
    return etree.XSLT(xslt_root)


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)


def set_outline_level(paragraph, level: int | None) -> None:
    if level is None:
        return
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def clear_outline_level(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None:
        p_pr.remove(outline)


def add_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break_before = p_pr.find(qn("w:pageBreakBefore"))
    if page_break_before is None:
        page_break_before = OxmlElement("w:pageBreakBefore")
        p_pr.append(page_break_before)


def apply_paragraph_style(paragraph, style: dict) -> None:
    source_style_name = style.get("source_style_name")
    if source_style_name:
        try:
            paragraph.style = paragraph.part.styles[source_style_name]
        except KeyError:
            pass

    if style.get("page_break_before"):
        add_page_break_before(paragraph)

    if style.get("clear_outline_level"):
        clear_outline_level(paragraph)
    set_outline_level(paragraph, style.get("outline_level"))

    if style.get("preserve_template_format"):
        return

    pf = paragraph.paragraph_format
    line_spacing_rule = style.get("line_spacing_rule", "exact")
    if line_spacing_rule == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing = None
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(style.get("line_spacing_pt", 20))
    pf.space_before = Pt(style.get("space_before_pt", 0))
    pf.space_after = Pt(style.get("space_after_pt", 0))
    pf.first_line_indent = Pt(style["first_line_indent_pt"]) if "first_line_indent_pt" in style else None
    pf.left_indent = Pt(style["left_indent_pt"]) if "left_indent_pt" in style else None

    alignment = style.get("alignment", "left")
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    east_asia = style.get("east_asia_font", "宋体")
    latin = style.get("latin_font", "Times New Roman")
    size_pt = style.get("size_pt", 10.5)
    bold = style.get("bold", False)
    for run in paragraph.runs:
        set_run_fonts(run, east_asia, latin, size_pt, bold=bold)


def apply_keywords_paragraph(paragraph, label: str, content: str, style: dict) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(style.get("line_spacing_pt", 20))
    pf.first_line_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    size_pt = style.get("size_pt", 10.5)
    label_run = paragraph.add_run(label)
    set_run_fonts(
        label_run,
        style.get("label_east_asia_font", "黑体"),
        style.get("label_latin_font", "Times New Roman"),
        size_pt,
        bold=True,
    )
    if content:
        content_run = paragraph.add_run(content if label.startswith("关键词") else f" {content}")
        set_run_fonts(
            content_run,
            style.get("content_east_asia_font", "宋体"),
            style.get("content_latin_font", "Times New Roman"),
            size_pt,
            bold=False,
        )


def normalize_heading_text(text: str, level: int) -> str:
    text = text.strip()
    if level == 1:
        text = re.sub(r"^第\s*[0-9一二三四五六七八九十百千]+\s*章[　\s]*", "", text)
    elif level in (2, 3, 4):
        text = re.sub(r"^\d+(?:\.\d+){0,3}[　\s]+", "", text)
    text = text.strip()
    return f"　{text}" if text else text


def normalize_heading_key(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", text).lower()


def format_back_matter_heading(text: str) -> str:
    compact = text.strip().replace(" ", "").replace("　", "")
    if compact == "致谢":
        return "致　　谢"
    if compact == "声明":
        return "声　　明"
    return compact


def normalized_template_candidates(text: str) -> list[str]:
    compact = text.strip().replace(" ", "").replace("　", "")
    candidates = [text.strip(), compact]
    if compact == "致谢":
        candidates.extend(["致　　谢", "致谢"])
    elif compact == "声明":
        candidates.extend(["声　　明", "声明"])
    elif compact == "目录":
        candidates.extend(["目　　录", "目录"])
    elif compact == "摘要":
        candidates.extend(["中文摘要", "摘要"])
    return candidates


def find_template_style_name(
    template_doc: Document | None, text: str, *, prefer_last: bool = False
) -> str | None:
    if template_doc is None:
        return None
    candidate_keys = {normalize_heading_key(item) for item in normalized_template_candidates(text)}
    matches: list[str] = []
    for paragraph in template_doc.paragraphs:
        if normalize_heading_key(paragraph.text.strip()) in candidate_keys and paragraph.style is not None:
            matches.append(paragraph.style.name)
    if matches:
        return matches[-1] if prefer_last else matches[0]
    return None


def build_special_heading_style(
    style_spec: dict,
    template_doc: Document | None,
    text: str,
    base_key: str,
    *,
    include_in_toc: bool = False,
    prefer_last_template_match: bool = False,
) -> dict:
    style = deepcopy(style_spec[base_key])
    template_style_name = find_template_style_name(
        template_doc,
        text,
        prefer_last=prefer_last_template_match,
    )
    if template_style_name:
        style["source_style_name"] = template_style_name
        style["preserve_template_format"] = True
    else:
        style.pop("source_style_name", None)
        style.pop("preserve_template_format", None)
    if include_in_toc:
        style["outline_level"] = 0
        style.pop("clear_outline_level", None)
    else:
        style["clear_outline_level"] = True
        style.pop("outline_level", None)
    return style


def classify_front_matter_heading(raw_text: str) -> str | None:
    key = normalize_heading_key(raw_text)
    if key in {"中文摘要", "摘要"}:
        return "cn_abstract"
    if key == "abstract":
        return "en_abstract"
    if key in {normalize_heading_key(v) for v in TOC_HEADINGS}:
        return "toc"
    return None


def split_markdown_sections(lines: list[str]) -> list[tuple[str | None, list[str]]]:
    sections: list[tuple[str | None, list[str]]] = []
    current_lines: list[str] = []
    current_tag: str | None = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            if current_lines:
                sections.append((current_tag, current_lines))
            raw_text = stripped[3:].strip()
            current_tag = classify_front_matter_heading(raw_text)
            current_lines = [line]
        else:
            current_lines.append(line)
    if current_lines:
        sections.append((current_tag, current_lines))
    return sections


def get_template_front_matter_order(template_doc: Document | None) -> list[str]:
    if template_doc is None:
        return list(FRONT_MATTER_ORDER_TAGS)
    order: list[str] = []
    for paragraph in template_doc.paragraphs:
        tag = classify_front_matter_heading(paragraph.text.strip())
        if tag and tag not in order:
            order.append(tag)
        if len(order) == len(FRONT_MATTER_ORDER_TAGS):
            break
    return order or list(FRONT_MATTER_ORDER_TAGS)


def reorder_front_matter_sections(
    lines: list[str], template_doc: Document | None
) -> tuple[list[str], bool]:
    sections = split_markdown_sections(lines)
    if template_doc is None:
        return lines, False

    front_order = get_template_front_matter_order(template_doc)
    include_cn_abstract_in_toc = (
        "toc" in front_order
        and "cn_abstract" in front_order
        and front_order.index("toc") < front_order.index("cn_abstract")
    )

    front_sections: dict[str, list[str]] = {}
    front_positions = [idx for idx, (tag, _) in enumerate(sections) if tag in FRONT_MATTER_ORDER_TAGS]
    if not front_positions:
        return lines, include_cn_abstract_in_toc

    original_front_order: list[str] = []
    for tag, section_lines in sections:
        if tag in FRONT_MATTER_ORDER_TAGS:
            front_sections[tag] = section_lines
            original_front_order.append(tag)

    ordered_tags = [tag for tag in front_order if tag in front_sections]
    ordered_tags.extend(tag for tag in original_front_order if tag not in ordered_tags)
    reordered_front_sections = [(tag, front_sections[tag]) for tag in ordered_tags]

    first_front_idx = front_positions[0]
    rebuilt_sections: list[tuple[str | None, list[str]]] = []
    inserted = False
    for idx, section in enumerate(sections):
        tag, section_lines = section
        if idx == first_front_idx and not inserted:
            rebuilt_sections.extend(reordered_front_sections)
            inserted = True
        if tag in FRONT_MATTER_ORDER_TAGS:
            continue
        rebuilt_sections.append((tag, section_lines))

    rebuilt_lines: list[str] = []
    for _, section_lines in rebuilt_sections:
        rebuilt_lines.extend(section_lines)
    return rebuilt_lines, include_cn_abstract_in_toc


def add_image(doc: Document, path: Path) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    section = doc.sections[-1]
    usable_width_emu = section.page_width - section.left_margin - section.right_margin
    usable_width_cm = max((usable_width_emu / 360000) - 0.6, 1)
    p.add_run().add_picture(str(path), width=Cm(min(14.0, usable_width_cm)))


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def apply_three_line_table(table) -> None:
    border_nil = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left=border_nil,
                right=border_nil,
                top=border_nil,
                bottom=border_nil,
            )

    if not table.rows:
        return

    top_border = {"val": "single", "sz": "8", "color": "000000"}
    mid_border = {"val": "single", "sz": "4", "color": "000000"}
    bottom_border = {"val": "single", "sz": "8", "color": "000000"}

    first_row = table.rows[0]
    last_row = table.rows[-1]
    for cell in first_row.cells:
        set_cell_border(cell, top=top_border, bottom=mid_border, left=border_nil, right=border_nil)
    for cell in last_row.cells:
        set_cell_border(cell, bottom=bottom_border, left=border_nil, right=border_nil)


def clear_table_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    cell.add_paragraph("")


def clone_template_table(doc: Document, template_doc: Document, table_index: int):
    template_table = template_doc.tables[table_index]
    tbl = deepcopy(template_table._tbl)
    doc._element.body.append(tbl)
    return doc.tables[-1]


def fill_table_from_rows(table, rows: list[list[str]], style: dict, math_transform=None) -> None:
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < len(rows):
        new_tr = deepcopy(table.rows[-1]._tr)
        table._tbl.append(new_tr)
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        for c_idx in range(len(tr.cells)):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = tr.cells[c_idx]
            clear_table_cell(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(style.get("line_spacing_pt", 20))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_with_inline_math(p, text, style, math_transform=math_transform)
            for run in p.runs:
                set_run_fonts(
                    run,
                    style.get("east_asia_font", "宋体"),
                    style.get("latin_font", "Times New Roman"),
                    style.get("size_pt", 10.5),
                    bold=style.get("bold", False),
                )


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for row in lines:
        cells = [c.strip().replace("`", "") for c in row.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return []
    return [rows[0]] + rows[2:]


def add_markdown_table(doc: Document, lines: list[str], style: dict, math_transform=None) -> None:
    rows = parse_markdown_table(lines)
    if not rows:
        return
    headers = rows[0]
    data_rows = rows[1:]
    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for r_idx, row in enumerate(data_rows, start=1):
        for c_idx, text in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = text
    if style.get("three_line_table", True):
        apply_three_line_table(table)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(style.get("line_spacing_pt", 20))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                raw_text = p.text
                for run in list(p.runs):
                    p._element.remove(run._r)
                add_text_with_inline_math(p, raw_text, style, math_transform=math_transform)
                for run in p.runs:
                    set_run_fonts(
                        run,
                        style.get("east_asia_font", "宋体"),
                        style.get("latin_font", "Times New Roman"),
                        style.get("size_pt", 10.5),
                        bold=style.get("bold", False),
                    )


def split_inline_math(text: str) -> list[tuple[str, str]]:
    parts: list[tuple[str, str]] = []
    pattern = re.compile(r"(?<!\\)\$(.+?)(?<!\\)\$")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            parts.append(("text", text[pos:match.start()]))
        parts.append(("math", match.group(1)))
        pos = match.end()
    if pos < len(text):
        parts.append(("text", text[pos:]))
    return parts if parts else [("text", text)]


def latex_to_omml(latex: str, math_transform) -> etree._Element | None:
    if not latex.strip() or math_transform is None or latex_to_mathml is None:
        return None
    try:
        mathml = latex_to_mathml(latex)
        mathml_root = etree.fromstring(mathml.encode("utf-8"))
        transformed = math_transform(mathml_root)
        root = transformed.getroot()
        return root
    except Exception:
        return None


def add_text_with_inline_math(paragraph, text: str, style: dict, math_transform=None) -> None:
    for part_type, value in split_inline_math(text):
        if part_type == "text":
            if value:
                run = paragraph.add_run(value)
                set_run_fonts(
                    run,
                    style.get("east_asia_font", "宋体"),
                    style.get("latin_font", "Times New Roman"),
                    style.get("size_pt", 10.5),
                    bold=style.get("bold", False),
                )
        else:
            omml = latex_to_omml(value, math_transform)
            if omml is None:
                run = paragraph.add_run(f"${value}$")
                set_run_fonts(
                    run,
                    style.get("east_asia_font", "宋体"),
                    style.get("latin_font", "Times New Roman"),
                    style.get("size_pt", 10.5),
                    bold=style.get("bold", False),
                )
            else:
                paragraph._p.append(etree.fromstring(etree.tostring(omml)))


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_toc_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    hint_run = OxmlElement("w:r")
    hint_text = OxmlElement("w:t")
    hint_text.text = "更新域后显示目录"
    hint_run.append(hint_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(hint_run)
    run._r.append(fld_end)


def insert_toc_block(
    doc: Document,
    style_spec: dict,
    heading_text: str = "目录",
    *,
    page_break_before: bool = False,
    template_doc: Document | None = None,
) -> None:
    if page_break_before:
        doc.add_page_break()
    heading = doc.add_paragraph()
    heading.add_run(heading_text)
    apply_paragraph_style(
        heading,
        build_special_heading_style(
            style_spec,
            template_doc,
            heading_text,
            "toc_heading",
        ),
    )
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    toc_paragraph.paragraph_format.line_spacing = Pt(20)
    add_toc_field(toc_paragraph)
    doc.add_page_break()


def create_document(template_path: Path | None, style_spec: dict, insert_toc: bool) -> Document:
    if template_path:
        doc = Document(str(template_path))
        enable_update_fields_on_open(doc)
        return doc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    enable_update_fields_on_open(doc)
    return doc


def resolve_replace_index(doc: Document, replace_from: str | None) -> int | None:
    markers = [replace_from] if replace_from else DEFAULT_REPLACE_MARKERS
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        for marker in markers:
            if marker and (text == marker or marker in text):
                return idx
    return None


def remove_body_from_index(doc: Document, paragraph_index: int | None) -> None:
    if paragraph_index is None:
        return
    target = doc.paragraphs[paragraph_index]._p
    body = doc._element.body
    remove = False
    for child in list(body.iterchildren()):
        if child == target:
            remove = True
        if remove and child.tag != qn("w:sectPr"):
            body.remove(child)


def build_doc(
    doc: Document,
    source: Path,
    image_map: dict[str, Path],
    style_spec: dict,
    style_payload: dict,
    skip_title: bool,
    should_insert_toc: bool,
    template_doc: Document | None,
) -> Document:
    lines, include_cn_abstract_in_toc = reorder_front_matter_sections(
        source.read_text(encoding="utf-8").splitlines(),
        template_doc,
    )
    current_section = ""
    in_code = False
    code_lang = ""
    pending_mermaid = False
    seen_content = False
    title_seen = False
    toc_inserted = False
    math_transform = build_math_transform()
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()
        if in_code:
            if stripped.startswith("```"):
                in_code = False
                if code_lang == "mermaid":
                    pending_mermaid = True
                code_lang = ""
            elif code_lang != "mermaid":
                p = doc.add_paragraph()
                p.add_run(lines[i].rstrip())
                apply_paragraph_style(
                    p,
                    {
                        "east_asia_font": "Consolas",
                        "latin_font": "Consolas",
                        "size_pt": 9,
                        "line_spacing_pt": 20,
                    },
                )
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped[3:].strip().lower()
            i += 1
            continue

        if stripped.startswith("$$"):
            block = [stripped[2:]]
            if stripped.endswith("$$") and len(stripped) > 4:
                block = [stripped[2:-2]]
            else:
                i += 1
                while i < len(lines):
                    current = lines[i].rstrip()
                    if current.strip().endswith("$$"):
                        block.append(current.strip()[:-2])
                        break
                    block.append(current)
                    i += 1
            equation = "\n".join(block).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            omml = latex_to_omml(equation, math_transform)
            if omml is None:
                run = p.add_run(f"$${equation}$$")
                set_run_fonts(run, "Times New Roman", "Times New Roman", 10.5, bold=False)
            else:
                p._p.append(etree.fromstring(etree.tostring(omml)))
            seen_content = True
            i += 1
            continue

        if stripped.startswith("# "):
            title_seen = True
            if not skip_title:
                p = doc.add_paragraph()
                p.add_run(stripped[2:].strip())
                apply_paragraph_style(p, style_spec["document_title"])
                seen_content = True
            if should_insert_toc and not toc_inserted and template_doc is None:
                insert_toc_block(doc, style_spec)
                toc_inserted = True
            i += 1
            continue

        if stripped.startswith("## "):
            raw_text = stripped[3:].strip()
            raw_key = normalize_heading_key(raw_text)
            if raw_key in {normalize_heading_key(v) for v in TOC_HEADINGS} and should_insert_toc:
                insert_toc_block(
                    doc,
                    style_spec,
                    heading_text=raw_text,
                    page_break_before=seen_content,
                    template_doc=template_doc,
                )
                toc_inserted = True
                current_section = raw_text
                seen_content = True
                i += 1
                continue
            if raw_key in {normalize_heading_key(v) for v in BACK_MATTER_HEADINGS}:
                text = format_back_matter_heading(raw_text)
            elif raw_text.startswith(APPENDIX_PREFIXES):
                text = normalize_heading_text(raw_text, 1)
            else:
                text = raw_text if raw_key in {normalize_heading_key(v) for v in SPECIAL_CENTERED_HEADINGS} else normalize_heading_text(raw_text, 1)
            p = doc.add_paragraph()
            p.add_run(text)
            if raw_key in {normalize_heading_key(v) for v in ABSTRACT_HEADINGS}:
                if raw_key == "abstract":
                    special_style = build_special_heading_style(
                        style_spec,
                        template_doc,
                        raw_text,
                        "abstract_heading_en",
                    )
                else:
                    special_style = build_special_heading_style(
                        {
                            **style_spec,
                            "abstract_heading_cn": style_spec.get(
                                "abstract_heading_cn",
                                style_spec["centered_heading"],
                            ),
                        },
                        template_doc,
                        raw_text,
                        "abstract_heading_cn",
                        include_in_toc=include_cn_abstract_in_toc,
                    )
                apply_paragraph_style(p, special_style)
            elif raw_key in {normalize_heading_key(v) for v in BACK_MATTER_HEADINGS}:
                apply_paragraph_style(
                    p,
                    build_special_heading_style(
                        style_spec,
                        template_doc,
                        text,
                        "centered_heading",
                        prefer_last_template_match=True,
                    ),
                )
            elif raw_text.startswith(APPENDIX_PREFIXES):
                apply_paragraph_style(p, style_spec.get("appendix_heading", style_spec["centered_heading"]))
            elif raw_key in {normalize_heading_key(v) for v in TOC_HEADINGS}:
                apply_paragraph_style(
                    p,
                    build_special_heading_style(
                        style_spec,
                        template_doc,
                        raw_text,
                        "toc_heading",
                    ),
                )
            elif raw_key in {normalize_heading_key(v) for v in SPECIAL_CENTERED_HEADINGS}:
                base_key = "abstract_heading_en" if raw_key == "abstract" else "centered_heading"
                apply_paragraph_style(
                    p,
                    build_special_heading_style(
                        style_spec,
                        template_doc,
                        raw_text,
                        base_key,
                    ),
                )
            else:
                apply_paragraph_style(p, style_spec["heading1"])
            current_section = text
            seen_content = True
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            p.add_run(normalize_heading_text(stripped[4:].strip(), 2))
            apply_paragraph_style(p, style_spec["heading2"])
            seen_content = True
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            p.add_run(normalize_heading_text(stripped[5:].strip(), 3))
            apply_paragraph_style(p, style_spec["heading3"])
            seen_content = True
            i += 1
            continue

        keyword_match = re.match(r"^(关键词[:：]|Keywords[:：])\s*(.*)$", stripped)
        if keyword_match:
            p = doc.add_paragraph()
            style_key = "keywords_cn" if keyword_match.group(1).startswith("关键词") else "keywords_en"
            apply_keywords_paragraph(p, keyword_match.group(1), keyword_match.group(2), style_spec[style_key])
            seen_content = True
            i += 1
            continue

        image_match = re.match(r"^\[此处插入截图：(.+?)\]$", stripped)
        if image_match:
            label = image_match.group(1).strip()
            image_path = image_map.get(label)
            if image_path and image_path.exists():
                add_image(doc, image_path)
                seen_content = True
            i += 1
            continue

        if re.match(r"^(图|表)\s*\d+(\.\d+)?", stripped):
            if pending_mermaid:
                image_path = image_map.get(stripped)
                if image_path and image_path.exists():
                    add_image(doc, image_path)
                pending_mermaid = False
            p = doc.add_paragraph()
            p.add_run(stripped.replace("`", ""))
            apply_paragraph_style(p, style_spec["caption"])
            seen_content = True
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = parse_markdown_table(table_lines)
            preferred_idx = style_payload.get("preferred_table_template_index")
            if template_doc is not None and preferred_idx is not None:
                try:
                    table = clone_template_table(doc, template_doc, preferred_idx)
                    fill_table_from_rows(table, rows, style_spec["table_text"], math_transform=math_transform)
                except Exception:
                    add_markdown_table(doc, table_lines, style_spec["table_text"], math_transform=math_transform)
            else:
                add_markdown_table(doc, table_lines, style_spec["table_text"], math_transform=math_transform)
            seen_content = True
            continue

        p = doc.add_paragraph()
        if current_section.replace(" ", "") == "参考文献" and re.match(r"^\[\d+\]", stripped):
            apply_paragraph_style(p, style_spec["reference"])
            add_text_with_inline_math(p, stripped.replace("`", ""), style_spec["reference"], math_transform=math_transform)
        elif current_section.lower() == "abstract":
            apply_paragraph_style(p, style_spec["abstract_body_en"])
            add_text_with_inline_math(p, stripped.replace("`", ""), style_spec["abstract_body_en"], math_transform=math_transform)
        else:
            apply_paragraph_style(p, style_spec["body"])
            add_text_with_inline_math(p, stripped.replace("`", ""), style_spec["body"], math_transform=math_transform)
        seen_content = True
        i += 1

    if should_insert_toc and not title_seen and not toc_inserted and template_doc is None:
        insert_toc_block(doc, style_spec)

    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("source")
    parser.add_argument("target")
    parser.add_argument("image_map", nargs="?")
    parser.add_argument("--template")
    parser.add_argument("--style-spec")
    parser.add_argument("--replace-from")
    parser.add_argument("--skip-title", action="store_true")
    parser.add_argument("--insert-toc", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source = Path(args.source)
    target = Path(args.target)
    image_map = load_image_map(Path(args.image_map)) if args.image_map else {}
    style_spec, style_payload = load_style_spec_payload(Path(args.style_spec)) if args.style_spec else (default_style_spec(), {})
    template_path = Path(args.template) if args.template else None

    target.parent.mkdir(parents=True, exist_ok=True)
    doc = create_document(template_path, style_spec, insert_toc=args.insert_toc)
    template_doc = Document(str(template_path)) if template_path else None
    if template_path:
        remove_body_from_index(doc, resolve_replace_index(doc, args.replace_from))

    build_doc(
        doc=doc,
        source=source,
        image_map=image_map,
        style_spec=style_spec,
        style_payload=style_payload,
        skip_title=args.skip_title or bool(template_path),
        should_insert_toc=args.insert_toc,
        template_doc=template_doc,
    )
    doc.save(target)
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
