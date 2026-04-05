from __future__ import annotations

import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ABSTRACT_CN_HEADINGS = {"摘要", "中文摘要"}
ABSTRACT_EN_HEADINGS = {"abstract"}
TOC_HEADINGS = {"目录"}
BACK_MATTER_HEADINGS = {"参考文献", "鸣谢", "致谢", "声明", "在学期间参加课题的研究成果"}


def length_score(text: str) -> int:
    return len(text.replace(" ", "").replace("\t", ""))


def first_non_empty_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return None


def paragraph_snapshot(index: int, paragraph) -> dict:
    text = paragraph.text.strip().replace("\t", " ")
    run = first_non_empty_run(paragraph)
    pf = paragraph.paragraph_format
    style = paragraph.style
    style_font = style.font if style else None
    style_pf = style.paragraph_format if style else None
    font_name = (run.font.name if run and run.font.name else None) or (style_font.name if style_font else None)
    size_pt = (run.font.size.pt if run and run.font.size else None) or (style_font.size.pt if style_font and style_font.size else None)
    bold = None
    if run and run.font.bold is not None:
        bold = run.font.bold
    elif style_font is not None:
        bold = style_font.bold
    return {
        "index": index,
        "text": text,
        "style_name": paragraph.style.name if paragraph.style else "",
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else None,
        "font_name": font_name,
        "size_pt": size_pt,
        "bold": bold,
        "italic": run.font.italic if run else None,
        "line_spacing": (
            pf.line_spacing.pt if getattr(pf.line_spacing, "pt", None) else pf.line_spacing
        ) or (
            style_pf.line_spacing.pt if style_pf and getattr(style_pf.line_spacing, "pt", None) else style_pf.line_spacing if style_pf else None
        ),
        "space_before_pt": (pf.space_before.pt if pf.space_before else None) or (style_pf.space_before.pt if style_pf and style_pf.space_before else None),
        "space_after_pt": (pf.space_after.pt if pf.space_after else None) or (style_pf.space_after.pt if style_pf and style_pf.space_after else None),
        "first_line_indent_pt": (pf.first_line_indent.pt if pf.first_line_indent else None) or (style_pf.first_line_indent.pt if style_pf and style_pf.first_line_indent else None),
        "left_indent_pt": (pf.left_indent.pt if pf.left_indent else None) or (style_pf.left_indent.pt if style_pf and style_pf.left_indent else None),
    }


def snapshot_to_style(snapshot: dict, default_alignment: str = "left") -> dict:
    style = {
        "source_style_name": snapshot.get("style_name") or None,
        "east_asia_font": snapshot.get("font_name") or "宋体",
        "latin_font": "Times New Roman",
        "size_pt": snapshot.get("size_pt") or 10.5,
        "bold": bool(snapshot.get("bold")),
        "alignment": default_alignment,
        "line_spacing_pt": snapshot.get("line_spacing") or 20,
        "space_before_pt": snapshot.get("space_before_pt") or 0,
        "space_after_pt": snapshot.get("space_after_pt") or 0,
    }
    if snapshot.get("first_line_indent_pt") is not None:
        style["first_line_indent_pt"] = snapshot["first_line_indent_pt"]
    if snapshot.get("left_indent_pt") is not None:
        style["left_indent_pt"] = snapshot["left_indent_pt"]
    return style


def style_name_to_spec(doc: Document, style_name: str, default_alignment: str = "left") -> dict | None:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    font = style.font
    pf = style.paragraph_format
    spec = {
        "source_style_name": style_name,
        "east_asia_font": font.name or "宋体",
        "latin_font": "Times New Roman",
        "size_pt": font.size.pt if font.size else 10.5,
        "bold": bool(font.bold),
        "alignment": default_alignment,
        "line_spacing_pt": (pf.line_spacing.pt if getattr(pf.line_spacing, "pt", None) else pf.line_spacing) or 20,
        "space_before_pt": pf.space_before.pt if pf.space_before else 0,
        "space_after_pt": pf.space_after.pt if pf.space_after else 0,
    }
    if pf.first_line_indent:
        spec["first_line_indent_pt"] = pf.first_line_indent.pt
    if pf.left_indent:
        spec["left_indent_pt"] = pf.left_indent.pt
    return spec


def classify_paragraph(snapshot: dict) -> str | None:
    text = snapshot["text"]
    style_name = (snapshot["style_name"] or "").lower()
    normalized = re.sub(r"[\s\u3000]+", "", text).lower()

    if not text:
        return None
    if style_name.startswith("toc"):
        return None
    if normalized in TOC_HEADINGS:
        return "toc_heading"
    if normalized in ABSTRACT_CN_HEADINGS:
        return "abstract_heading_cn"
    if normalized in {"设计总说明"} | BACK_MATTER_HEADINGS:
        return "centered_heading"
    if normalized in ABSTRACT_EN_HEADINGS | {"introduction"}:
        return "abstract_heading_en"
    if text.startswith("关键词") or text.startswith("Keywords"):
        return "keywords_cn" if text.startswith("关键词") else "keywords_en"
    if re.match(r"^(图|表)\s*\d+(\.\d+)?", text) or "caption" in style_name:
        return "caption"
    if re.match(r"^\[\d+\]", text):
        return "reference"
    if style_name.startswith("heading 4"):
        return "heading4"
    if style_name.startswith("heading 3"):
        return "heading3"
    if style_name.startswith("heading 2"):
        return "heading2"
    if style_name.startswith("heading 1"):
        return "heading1"
    if style_name.startswith("附录标题") or normalized.startswith("附录"):
        return "appendix_heading"
    if "title" in style_name and length_score(text) <= 40:
        return "document_title"
    if re.match(r"^\d+\.\d+\.\d+\.\d+", text):
        return "heading4"
    if re.match(r"^\d+\.\d+\.\d+", text):
        return "heading3"
    if re.match(r"^\d+\.\d+", text):
        return "heading2"
    if re.match(r"^\d+\s", text):
        return "heading1"
    if length_score(text) > 25:
        return "body"
    return None


def border_val(element, edge: str) -> str | None:
    tc_pr = element._tc.tcPr
    if tc_pr is None:
        return None
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        return None
    border = tc_borders.find(qn(f"w:{edge}"))
    if border is None:
        return None
    return border.get(qn("w:val"))


def analyze_table(table, index: int) -> dict:
    has_vertical = False
    has_top = False
    has_bottom = False
    has_mid = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            left = border_val(cell, "left")
            right = border_val(cell, "right")
            top = border_val(cell, "top")
            bottom = border_val(cell, "bottom")
            if left not in {None, "nil"} or right not in {None, "nil"}:
                has_vertical = True
            if row_idx == 0 and top not in {None, "nil"}:
                has_top = True
            if row_idx == 0 and bottom not in {None, "nil"}:
                has_mid = True
            if row_idx == len(table.rows) - 1 and bottom not in {None, "nil"}:
                has_bottom = True
    first_row = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
    return {
        "index": index,
        "rows": len(table.rows),
        "cols": len(table.columns),
        "first_row": first_row,
        "has_vertical_borders": has_vertical,
        "has_top_border": has_top,
        "has_header_separator": has_mid,
        "has_bottom_border": has_bottom,
        "looks_like_three_line_table": (not has_vertical) and has_top and has_bottom,
    }


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python analyze_docx_styles.py <template-or-sample.docx> [output.json]")
        return 1

    path = Path(sys.argv[1])
    doc = Document(str(path))

    paragraphs = []
    classified = {}
    counts = Counter()
    first_body_heading = None
    first_toc = None
    last_toc = None
    first_abstract_cn = None
    first_abstract_en = None
    first_back_matter = None
    first_appendix = None

    for idx, paragraph in enumerate(doc.paragraphs):
        snap = paragraph_snapshot(idx, paragraph)
        if not snap["text"]:
            continue
        paragraphs.append(snap)
        kind = classify_paragraph(snap)
        if kind:
            counts[kind] += 1
            classified.setdefault(kind, []).append(snap)
        if first_toc is None and kind == "toc_heading":
            first_toc = snap
        if kind == "abstract_heading_cn" and first_abstract_cn is None:
            first_abstract_cn = snap
        if kind == "abstract_heading_en" and first_abstract_en is None:
            first_abstract_en = snap
        if kind == "centered_heading" and first_back_matter is None:
            first_back_matter = snap
        if kind == "appendix_heading" and first_appendix is None:
            first_appendix = snap
        if kind and kind.startswith("heading") and first_body_heading is None:
            first_body_heading = snap
        if (snap["style_name"] or "").lower().startswith("toc"):
            last_toc = snap

    replace_from = None
    if first_abstract_cn:
        replace_from = first_abstract_cn
    elif first_body_heading:
        replace_from = first_body_heading

    style_spec = {}
    preferred_styles = {
        "heading1": ("Heading 1", "center"),
        "heading2": ("Heading 2", "left"),
        "heading3": ("Heading 3", "left"),
        "heading4": ("Heading 4", "left"),
        "body": ("Normal", "left"),
        "appendix_heading": ("附录标题 1", "center"),
    }

    for key, (style_name, alignment) in preferred_styles.items():
        spec = style_name_to_spec(doc, style_name, default_alignment=alignment)
        if spec:
            style_spec[key] = spec

    for key in [
        "document_title",
        "centered_heading",
        "abstract_heading_cn",
        "abstract_heading_en",
        "appendix_heading",
        "heading1",
        "heading2",
        "heading3",
        "heading4",
        "body",
        "caption",
        "reference",
        "toc_heading",
    ]:
        if key in style_spec:
            continue
        if key in classified:
            snap = classified[key][0]
            default_alignment = "center" if key in {"document_title", "centered_heading", "abstract_heading_cn", "abstract_heading_en", "toc_heading", "caption", "appendix_heading"} else "left"
            style_spec[key] = snapshot_to_style(snap, default_alignment=default_alignment)

    if "body" in classified:
        style_spec["table_text"] = snapshot_to_style(classified["body"][0], default_alignment="center")
        style_spec["abstract_body_en"] = {
            **snapshot_to_style(classified["body"][0], default_alignment="left"),
            "east_asia_font": "Times New Roman",
        }

    if "keywords_cn" in classified:
        snap = classified["keywords_cn"][0]
        style_spec["keywords_cn"] = {
            "source_style_name": snap.get("style_name") or None,
            "label_east_asia_font": snap.get("font_name") or "黑体",
            "label_latin_font": "Times New Roman",
            "content_east_asia_font": "宋体",
            "content_latin_font": "Times New Roman",
            "size_pt": snap.get("size_pt") or 10.5,
            "alignment": "left",
            "line_spacing_pt": snap.get("line_spacing") or 20,
        }

    if "keywords_en" in classified:
        snap = classified["keywords_en"][0]
        style_spec["keywords_en"] = {
            "source_style_name": snap.get("style_name") or None,
            "label_east_asia_font": "Times New Roman",
            "label_latin_font": "Times New Roman",
            "content_east_asia_font": "Times New Roman",
            "content_latin_font": "Times New Roman",
            "size_pt": snap.get("size_pt") or 10.5,
            "alignment": "left",
            "line_spacing_pt": snap.get("line_spacing") or 20,
        }

    result = {
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "detected_counts": dict(counts),
        "markers": {
            "first_abstract_cn_text": first_abstract_cn["text"] if first_abstract_cn else None,
            "first_abstract_en_text": first_abstract_en["text"] if first_abstract_en else None,
            "first_toc_text": first_toc["text"] if first_toc else None,
            "first_body_heading_text": first_body_heading["text"] if first_body_heading else None,
            "first_back_matter_text": first_back_matter["text"] if first_back_matter else None,
            "first_appendix_text": first_appendix["text"] if first_appendix else None,
            "last_toc_text": last_toc["text"] if last_toc else None,
            "replace_from_text": replace_from["text"] if replace_from else None,
        },
        "style_spec": style_spec,
        "examples": {
            key: [item["text"] for item in values[:5]]
            for key, values in classified.items()
        },
        "table_templates": [analyze_table(table, idx) for idx, table in enumerate(doc.tables)],
    }

    three_line = [t for t in result["table_templates"] if t["looks_like_three_line_table"]]
    result["preferred_table_template_index"] = three_line[0]["index"] if three_line else None

    payload = json.dumps(result, ensure_ascii=False, indent=2)
    if len(sys.argv) >= 3:
        target = Path(sys.argv[2])
        target.write_text(payload, encoding="utf-8")
        print(target)
    else:
        print(payload)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
